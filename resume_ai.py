import os
import sys
import json
import re
import shutil

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.document import Document as _Document
from openai import OpenAI


MODEL = "gpt-5.6-luna"


# ============================================================
# BASIC HELPERS
# ============================================================

def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_json(path, data):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def get_client():
    api_key = os.environ.get("OPENAI_API_KEY")

    if not api_key:
        raise RuntimeError(
            "OPENAI_API_KEY environment variable is missing."
        )

    return OpenAI(api_key=api_key)


# ============================================================
# DOCX PARAGRAPH HELPERS
# ============================================================

def iter_cell_paragraphs(cell, table_path):
    """
    Recursively return all paragraphs inside a table cell,
    including paragraphs inside nested tables.
    """

    results = []

    # Direct paragraphs inside this cell
    for paragraph_index, paragraph in enumerate(cell.paragraphs):

        results.append({
            "paragraph_id": (
                f"{table_path}:p:{paragraph_index}"
            ),
            "location": {
                "type": "table_cell",
                "table_path": table_path,
                "paragraph_index": paragraph_index
            },
            "paragraph": paragraph
        })

    # Nested tables
    for nested_table_index, nested_table in enumerate(cell.tables):

        nested_table_path = (
            f"{table_path}:nested_table:{nested_table_index}"
        )

        for row_index, row in enumerate(nested_table.rows):

            for cell_index, nested_cell in enumerate(row.cells):

                nested_cell_path = (
                    f"{nested_table_path}:r:{row_index}:c:{cell_index}"
                )

                results.extend(
                    iter_cell_paragraphs(
                        nested_cell,
                        nested_cell_path
                    )
                )

    return results


def get_all_paragraphs(doc):

    results = []

    # Normal paragraphs
    for paragraph_index, paragraph in enumerate(doc.paragraphs):

        results.append({
            "paragraph_id": f"body:p:{paragraph_index}",
            "location": {
                "type": "body",
                "paragraph_index": paragraph_index
            },
            "paragraph": paragraph
        })

    # Normal tables
    for table_index, table in enumerate(doc.tables):

        table_path = f"table:{table_index}"

        for row_index, row in enumerate(table.rows):

            for cell_index, cell in enumerate(row.cells):

                cell_path = (
                    f"{table_path}:r:{row_index}:c:{cell_index}"
                )

                results.extend(
                    iter_cell_paragraphs(
                        cell,
                        cell_path
                    )
                )

    # XML tables
    for table_index, tbl in enumerate(
        doc._element.iter()
    ):

        if not tbl.tag.endswith("}tbl"):
            continue

        table_path = f"xml_table:{table_index}"

        for row_index, tr in enumerate(tbl):

            if not tr.tag.endswith("}tr"):
                continue

            for cell_index, tc in enumerate(tr):

                if not tc.tag.endswith("}tc"):
                    continue

                cell_path = (
                    f"{table_path}:r:{row_index}:c:{cell_index}"
                )

                paragraph_index = 0

                for p_element in tc.iter():

                    if not p_element.tag.endswith("}p"):
                        continue

                    paragraph = Paragraph(
                        p_element,
                        doc
                    )

                    results.append({
                        "paragraph_id": (
                            f"{cell_path}:p:{paragraph_index}"
                        ),
                        "location": {
                            "type": "table_cell",
                            "table_path": table_path,
                            "row_index": row_index,
                            "cell_index": cell_index,
                            "paragraph_index": paragraph_index
                        },
                        "paragraph": paragraph
                    })

                    paragraph_index += 1

    return results
    
def build_paragraph_map(doc):
    """
    Build:

        paragraph_id -> actual python-docx Paragraph

    This allows exact editing of paragraphs inside table cells.
    """

    paragraph_map = {}

    for item in get_all_paragraphs(doc):

        paragraph_map[
            item["paragraph_id"]
        ] = item["paragraph"]

    return paragraph_map


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx(docx_path, output_path):

    doc = Document(docx_path)

    all_paragraphs = get_all_paragraphs(doc)

    paragraphs = []

    for item in all_paragraphs:

        paragraph = item["paragraph"]

        text = paragraph.text.strip()

        if not text:
            continue

        location = item["location"]

        paragraphs.append({
            "paragraph_id": item["paragraph_id"],
            "location": location,
            "text": text,
            "style": (
                paragraph.style.name
                if paragraph.style
                else ""
            ),
            "alignment": str(paragraph.alignment),
            "is_table_paragraph": (
                location["type"] == "table_cell"
            )
        })

    data = {
        "document": {
            "paragraph_count": len(paragraphs)
        },
        "paragraphs": paragraphs
    }

    save_json(output_path, data)

    print(
        f"Extracted {len(paragraphs)} editable paragraphs "
        f"(including table/cell paragraphs)."
    )

    table_count = sum(
        1
        for item in paragraphs
        if item.get("is_table_paragraph")
    )

    print(
        f"Table/cell paragraphs detected: {table_count}"
    )


# ============================================================
# OPENAI JSON EXTRACTION
# ============================================================

def get_response_text(response):

    if hasattr(response, "output_text"):
        return response.output_text

    return str(response)


def clean_json_text(text):

    text = text.strip()

    if text.startswith("```"):

        text = re.sub(
            r"^```(?:json)?",
            "",
            text
        )

        text = re.sub(
            r"```$",
            "",
            text
        )

        text = text.strip()

    return text


# ============================================================
# ATS ANALYSIS
# ============================================================

def ats_analysis(
    resume_json_path,
    jd_path,
    output_path
):

    resume = load_json(resume_json_path)

    with open(
        jd_path,
        "r",
        encoding="utf-8"
    ) as f:

        jd = f.read()

    prompt = f"""
You are an ATS resume evaluation system.

Compare the resume against the job description.

IMPORTANT:

The resume is real candidate data.

Evaluate the resume exactly as written.

Consider:

- Skills section
- Experience section
- Projects section
- Certifications
- Keywords
- Responsibilities
- Cloud platforms
- DevOps tools
- CI/CD
- Infrastructure as Code
- Monitoring
- Security
- Scripting

Do not assume experience that is not present.

Return ONLY valid JSON.

Required JSON:

{{
  "ats_score": 0,
  "matched_skills": [],
  "missing_skills": [],
  "suggestions": []
}}

ATS_SCORE must be an integer from 0 to 100.

JOB DESCRIPTION:

{jd}

RESUME:

{json.dumps(
    resume,
    indent=2,
    ensure_ascii=False
)}
"""

    client = get_client()

    response = client.responses.create(
        model=MODEL,
        input=prompt
    )

    text = clean_json_text(
        get_response_text(response)
    )

    result = json.loads(text)

    if "ats_score" not in result:

        raise RuntimeError(
            "AI response does not contain ats_score."
        )

    score = int(result["ats_score"])

    if score < 0 or score > 100:

        raise RuntimeError(
            "Invalid ATS score returned by AI."
        )

    result["ats_score"] = score

    save_json(
        output_path,
        result
    )

    print(
        json.dumps(
            result,
            indent=2,
            ensure_ascii=False
        )
    )


# ============================================================
# RESUME REWRITE
# ============================================================

def rewrite_resume(
    original_docx,
    resume_json_path,
    jd_path,
    output_path
):

    resume = load_json(resume_json_path)

    with open(
        jd_path,
        "r",
        encoding="utf-8"
    ) as f:

        jd = f.read()

    prompt = f"""
You are a professional ATS resume writer.

Your job is to improve ONLY the wording, ordering,
and JD alignment of:

1. Professional Summary
2. Skills
3. Experience
4. Project Titles
5. Project Bullet Points

The original DOCX structure is the source of truth.

The RESUME STRUCTURE JSON contains every editable
paragraph in the document, including paragraphs inside
tables and table cells.

IMPORTANT:

A paragraph can be located in:

- the main document body
- a table
- a sidebar implemented using a table cell
- a nested table

Every paragraph has a unique "paragraph_id".

Examples:

body:p:0

body:p:5

table:0:r:0:c:1:p:0

table:0:r:0:c:1:p:2

You MUST use the exact paragraph_id from the
RESUME STRUCTURE JSON.

NEVER invent a paragraph_id.

NEVER convert paragraph_id into a numeric index.

NEVER use "paragraph_index" as the replacement identifier.

============================================================
STRICT FACTUAL RULES
============================================================

DO NOT change:

- Candidate name
- Contact information
- Education
- Degree
- University
- Company names
- Job titles
- Employment dates
- Actual experience facts
- Project facts
- Certifications
- Factual information

DO NOT invent:

- Skills
- Tools
- Technologies
- Achievements
- Metrics
- Companies
- Projects
- Responsibilities

============================================================
MANDATORY REWRITING
============================================================

Rewrite only when the change improves ATS alignment, keyword relevance,
clarity, or impact.

Professional Summary:
Rewrite when JD alignment can be improved.

Skills:
Rewrite when supported JD-relevant skills can be reordered or added.

Experience:
Rewrite when bullets can be better aligned with the JD without changing facts.

Project Titles:
Rewrite only when the existing title can be made more relevant without
changing the actual project.

Project Bullet Points:
Rewrite when JD-relevant wording can be improved without changing facts.

Do NOT change text merely for the sake of changing it.

Even if a section already looks good,
rewrite it to better align with the target JD.

============================================================
SKILLS SECTION — CRITICAL
============================================================

============================================================
SKILLS IDENTIFICATION — MANDATORY
============================================================

Before generating replacements, inspect EVERY paragraph in
RESUME STRUCTURE and identify the actual Skills content.

You MUST return at least ONE replacement with:

"section": "Skills"

Use the EXACT paragraph_id of the existing Skills content.

NEVER assume Skills is body:p:XX.

NEVER use a paragraph just because its number looks correct.

If Skills is inside a table or sidebar, preserve its exact
table/cell paragraph_id.

If Skills contains multiple paragraphs, return one replacement
for each existing Skills content paragraph.

Existing Skills category headings MAY be rewritten for better
JD/ATS alignment, but ONLY if useful.

If a Skills heading is rewritten:

- use its exact existing paragraph_id
- keep it in the exact same location
- keep the same formatting
- do not create a new heading
- do not delete a heading
- do not create a new category
- new heading text MUST NOT be longer than the original heading

The Skills section MUST NOT be left unchanged.

You MUST:

1. Extract ALL skills and technologies supported by evidence
   anywhere in the resume.

2. Compare supported skills against the target JD.

3. Reorder supported skills according to JD relevance.

4. Put the most JD-relevant supported skills first.

5. Add supported skills that are missing from the current
   Skills section when those skills are explicitly supported
   elsewhere in the resume.

6. Evidence may come from:

   - Professional Summary
   - Experience
   - Projects
   - Existing Skills
   - Certifications

7. Examples:

   - Helm if used in Projects
   - Bash if used in Projects
   - Python if used in Projects
   - Loki if used in Projects
   - SonarQube if used in Projects
   - Trivy if used in Projects
   - CI/CD if Jenkins or GitHub Actions pipelines are present
   - Infrastructure as Code if Terraform is present
   - Container Orchestration if Docker/Kubernetes/EKS/ECS are present
   - Monitoring & Observability if Prometheus/Grafana/CloudWatch are present

8. NEVER add a technology just because it appears in the JD.

9. NEVER invent unsupported skills.

10. Preserve the existing Skills categories where possible.

11. Reorder skills inside the existing categories based on
    JD relevance.

12. Add supported skills to the most appropriate existing category.

13. Do NOT create a new Skills section.

14. Do NOT move Skills into another section.

15. Do NOT place Skills content into a Professional Summary,
    Experience, Project, or other paragraph.

============================================================
TABLE / SIDEBAR SAFETY — CRITICAL
============================================================

The original resume may use a table to create a two-column layout.

The right-side Skills section may exist inside a table cell.

Therefore:

- A Skills paragraph inside a table cell MUST be updated
  at its original table/cell location.

- NEVER place Skills text into a body paragraph merely because
  that body paragraph has a convenient paragraph index.

- NEVER move Skills content from a table cell to the body.

- NEVER copy Skills into another paragraph.

- NEVER duplicate Skills.

- If the Skills section contains multiple paragraphs,
  rewrite those paragraphs individually while preserving
  their original paragraph structure.

- Do NOT combine multiple Skills paragraphs into one paragraph.

- Do NOT create additional Skills paragraphs.

- Return one replacement for each existing Skills paragraph
  that actually needs rewriting.

============================================================
SECTION IDENTIFICATION
============================================================

Use the text and location information in RESUME STRUCTURE
to determine which paragraphs belong to:

- Professional Summary
- Skills
- Experience
- Project Titles
- Project Bullet Points

Pay special attention to table/cell locations.

If the Skills heading and Skills content are inside a table,
the replacement paragraphs MUST use those same table/cell
paragraph_ids.

The Skills replacement MUST NEVER target a random body paragraph.

# ============================================================
# LENGTH / LAYOUT — ABSOLUTE
# ============================================================

The ORIGINAL resume layout is fixed.

The rewritten content MUST fit inside the EXACT SAME
physical space as the original content.

For EVERY replacement:

- Keep the rewritten content close to the original physical length.
- Do NOT create additional bullets.
- Do NOT create new paragraphs.
- Do NOT change document structure.
- Do NOT intentionally make content substantially longer.
- Prefer concise ATS-focused wording.
- Preserve the original layout as much as possible.

If more JD keywords are needed, PRIORITIZE the most relevant
supported keywords and remove/reorder lower-priority wording.

NEVER solve keyword matching by making the section longer.

This rule applies to:

- Professional Summary
- Skills
- Experience
- Project Titles
- Project Bullet Points

The goal is:

SAME SPACE + BETTER JD ALIGNMENT.

NOT:

MORE SPACE + MORE KEYWORDS.

For Skills specifically:

- Preserve every existing Skills category.
- Keep Version Control in its original location.
- Do not allow any Skills category to move to another page.
- Do not create new Skills categories.
- Do not create new Skills paragraphs.

- Keep Skills paragraphs concise.
- Do not substantially increase the original Skills paragraph length.
- Prioritize JD-relevant supported skills.
- Remove lower-value or redundant wording when necessary.
- Never add unsupported skills.

- NEVER make a Skills paragraph longer than its original text.

- Reorder existing supported skills according to JD relevance.

- Add supported skills only when they can fit by removing,
  replacing, or reordering lower-priority skills.

- Do NOT simply append new skills to the end of a Skills
  paragraph.

- Preserve important existing technology names and their
  specific services where possible.

- Do NOT remove important technology/service names merely
  to add generic keywords.

- If a JD keyword cannot fit within the existing character
  limit, skip that keyword rather than increasing the length.

- Keep each Skills category in the exact same paragraph and
  exact same location.

- Keep the total Skills section within the original available
  physical space.

- Do NOT duplicate Skills content.

- Do NOT move Skills content to another paragraph or section.

Formatting must remain unchanged:

- same font
- same font size
- same bold/underline
- same alignment
- same spacing
- same table
- same cell
- same column
- same margins
- same location

============================================================
OUTPUT FORMAT — VERY IMPORTANT
============================================================

Return ONLY valid JSON.

Use this exact format:

{{
  "replacements": [
    {{
      "paragraph_id": "body:p:5",
      "section": "Professional Summary",
      "original_text": "existing paragraph text",
      "new_text": "rewritten paragraph text"
    }},
    {{
      "paragraph_id": "table:0:r:0:c:1:p:2",
      "section": "Skills",
      "original_text": "existing skills paragraph text",
      "new_text": "rewritten skills paragraph text"
    }}
  ]
}}

============================================================
MANDATORY OUTPUT REQUIREMENTS
============================================================

The "replacements" array MUST contain at least one replacement
for each of:

- Professional Summary
- Skills
- Experience
- Project Titles
- Project Bullet Points

For Skills:

- Use the actual Skills paragraph_id(s) from RESUME STRUCTURE.
- Preserve the original table/cell location.
- Return the rewritten Skills text in those same paragraphs.
- Do not duplicate Skills elsewhere.

Every replacement MUST contain:

- paragraph_id
- section
- original_text
- new_text

Every paragraph_id MUST exist in RESUME STRUCTURE.

NEVER return:

"paragraph_index"

NEVER return a numeric-only paragraph identifier.

============================================================
FORMATTING PRESERVATION RULES
============================================================

- Do not create new sections.
- Do not create new tables.
- Do not create new headings.
- Do not change document structure.
- Do not move paragraphs.
- Do not move content between columns.
- Do not move content between table cells.
- Do not duplicate content.
- Only replace existing paragraph text.
- Preserve the original paragraph locations.
- Assume original font, size, spacing, alignment,
  table structure, and formatting will be preserved.

============================================================
JOB DESCRIPTION
============================================================

{jd}

============================================================
RESUME STRUCTURE
============================================================

{json.dumps(
    resume,
    indent=2,
    ensure_ascii=False
)}
"""

    client = get_client()

    response = client.responses.create(
        model=MODEL,
        input=prompt
    )

    text = clean_json_text(
        get_response_text(response)
    )

    result = json.loads(text)

    if "replacements" not in result:
        raise RuntimeError(
            "AI response does not contain replacements."
        )

    if not isinstance(
        result["replacements"],
        list
    ):
        raise RuntimeError(
            "AI replacements must be a list."
        )

    save_json(
        output_path,
        result
    )

    print(
        json.dumps(
            result,
            indent=2,
            ensure_ascii=False
        )
    )


# ============================================================
# REPLACE PARAGRAPH TEXT
# ============================================================

def replace_paragraph_text(
    paragraph,
    new_text
):

    runs = paragraph.runs

    if not runs:

        paragraph.add_run(new_text)

        return

    # --------------------------------------------------------
    # Preserve the first run's character formatting.
    # --------------------------------------------------------

def replace_paragraph_text(paragraph, new_text):

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    first_run = paragraph.runs[0]

    first_run.text = new_text

    for run in paragraph.runs[1:]:
        run.text = ""


# ============================================================
# LENGTH SAFETY CHECK
# ============================================================

# Skills fitting function

def fit_skills_to_original_length(
    original_text,
    new_text
):
    max_len = len(
        re.sub(
            r"\s+",
            " ",
            str(original_text)
        ).strip()
    )

    new_text = re.sub(
        r"\s+",
        " ",
        str(new_text)
    ).strip()

    if len(new_text) <= max_len:
        return new_text

    skills = [
        skill.strip()
        for skill in new_text.split(",")
        if skill.strip()
    ]

    selected = []

    for skill in skills:

        candidate = ", ".join(
            selected + [skill]
        )

        if len(candidate) <= max_len:
            selected.append(skill)

    return ", ".join(selected)

# Existing function

def validate_replacement_length(
    original_text,
    new_text,
    paragraph_id,
    section
):
    """
    Layout safety:
    Skills must not become longer.
    Other sections can increase by up to 20%.
    """

    original_len = len(
        re.sub(r"\s+", " ", str(original_text)).strip()
    )

    new_len = len(
        re.sub(r"\s+", " ", str(new_text)).strip()
    )

    if section == "Skills":
        max_len = original_len
    else:
        max_len = int(original_len * 1.20)
        
    if new_len > max_len:
        raise RuntimeError(
            "Layout safety violation.\n"
            f"Paragraph ID: {paragraph_id}\n"
            f"Section: {section}\n"
            f"Original length: {original_len}\n"
            f"New length: {new_len}\n"
            f"Maximum allowed length: {max_len}"
        )


# ============================================================
# APPLY AI REWRITE
# ============================================================

def apply_rewrite(
    original_docx,
    rewritten_json,
    output_docx
):

    doc = Document(original_docx)

    rewritten = load_json(
        rewritten_json
    )

    replacements = rewritten.get(
        "replacements",
        []
    )

    if not replacements:

        raise RuntimeError(
            "AI returned no replacements."
        )

    # --------------------------------------------------------
    # Build exact paragraph map.
    #
    # This includes table/cell paragraphs.
    # --------------------------------------------------------

    paragraph_map = build_paragraph_map(doc)

    print(
        f"Editable paragraph locations found: "
        f"{len(paragraph_map)}"
    )

    # --------------------------------------------------------
    # Mandatory sections
    # --------------------------------------------------------

    required_sections = {
        "Professional Summary",
        "Skills",
        "Experience",
        "Project Titles",
        "Project Bullet Points"
    }

    returned_sections = {
        item.get("section")
        for item in replacements
        if item.get("section")
    }

    missing_sections = (
        required_sections -
        returned_sections
    )

    if missing_sections:

        raise RuntimeError(
            "AI rewrite is missing mandatory sections: "
            + ", ".join(
                sorted(missing_sections)
            )
        )

    print(
        "AI returned all mandatory resume sections."
    )

    # --------------------------------------------------------
    # Duplicate protection
    # --------------------------------------------------------

    used_paragraph_ids = set()

    # Track where Skills replacements are being applied.
    skills_locations = []

    # --------------------------------------------------------
    # Apply replacements
    # --------------------------------------------------------

    for item in replacements:

        paragraph_id = item.get(
            "paragraph_id"
        )

        section = item.get(
            "section"
        )

        original_text = item.get(
            "original_text",
            ""
        )

        new_text = item.get(
            "new_text"
        )

        # ----------------------------------------------------
        # Validate fields
        # ----------------------------------------------------

        if not paragraph_id:

            raise RuntimeError(
                "AI replacement is missing paragraph_id."
            )

        if not section:

            raise RuntimeError(
                f"Replacement {paragraph_id} "
                f"is missing section."
            )

        if new_text is None:

            raise RuntimeError(
                f"Replacement {paragraph_id} "
                f"is missing new_text."
            )

    # --------------------------------------------------------
    # HARD LENGTH / LAYOUT SAFETY CHECK
    # --------------------------------------------------------

    if section == "Skills":

        new_text = fit_skills_to_original_length(
            original_text,
            str(new_text)
        )

    else:

        validate_replacement_length(
            original_text,
            str(new_text),
            paragraph_id,
            section
        )

    # --------------------------------------------------------
    # Reject old unsafe format
    # --------------------------------------------------------

    if "paragraph_index" in item:

        raise RuntimeError(
            "Unsafe AI output detected: "
            f"paragraph_index used for {paragraph_id}. "
            "AI must use paragraph_id."
        )

    # --------------------------------------------------------
    # Prevent duplicate replacement of same paragraph
    # --------------------------------------------------------

    if paragraph_id in used_paragraph_ids:

        raise RuntimeError(
            "Duplicate replacement detected for "
            f"paragraph_id: {paragraph_id}"
        )

    used_paragraph_ids.add(
        paragraph_id
    )

    # --------------------------------------------------------
    # Exact paragraph existence check
    # --------------------------------------------------------

    if paragraph_id not in paragraph_map:

        raise RuntimeError(
            "AI returned an invalid paragraph_id: "
            f"{paragraph_id}"
        )

    paragraph = paragraph_map[
        paragraph_id
    ]

    actual_text = paragraph.text.strip()

    # --------------------------------------------------------
    # Validate original text when provided
    # --------------------------------------------------------

    if original_text:

        normalized_actual = re.sub(
            r"\s+",
            " ",
            actual_text
        ).strip()

        normalized_original = re.sub(
            r"\s+",
            " ",
            str(original_text)
        ).strip()

        if normalized_actual != normalized_original:

            raise RuntimeError(
                "Paragraph safety validation failed.\n"
                f"Paragraph ID: {paragraph_id}\n"
                f"Expected: {normalized_original[:200]}\n"
                f"Actual:   {normalized_actual[:200]}"
            )

    # ----------------------------------------------------
    # Apply exact paragraph replacement
    # ----------------------------------------------------

    replace_paragraph_text(
        paragraph,
        str(new_text)
    )

    if section == "Skills":
        skills_locations.append(
            paragraph_id
        )

    # --------------------------------------------------------
    # Logging
    # --------------------------------------------------------

    print(
        "Updating:"
    )

    print(
        f"  paragraph_id = {paragraph_id}"
    )

    print(
        f"  section      = {section}"
    )

    print(
        f"  original     = {actual_text[:100]}"
    )

    print(
        f"  new          = {str(new_text)[:100]}"
    )

    # --------------------------------------------------------
    # Additional Skills safety check
    # --------------------------------------------------------

    if not skills_locations:

        raise RuntimeError(
            "No Skills paragraph was actually updated."
        )

    print(
        "Skills paragraphs updated at original locations:"
    )

    for paragraph_id in skills_locations:

        location = paragraph_map[
            paragraph_id
        ]

        print(
            f"  - {paragraph_id}: "
            f"{location.text[:80]}"
        )

    # --------------------------------------------------------
    # Save output
    # --------------------------------------------------------

    doc.save(
        output_docx
    )

    print(
        f"Created: {output_docx}"
    )


# ============================================================
# OPTIMIZE RESUME
# ============================================================

def optimize_resume(
    original_docx,
    jd_path,
    target_score=80,
    max_attempts=3
):

    current_docx = original_docx

    best_score = 0
    best_resume = original_docx

    for attempt in range(
        1,
        max_attempts + 1
    ):

        print(
            "\n"
            + "=" * 60
        )

        print(
            f"ATTEMPT {attempt}"
        )

        print(
            "=" * 60
        )

        # ----------------------------------------------------
        # Extract
        # ----------------------------------------------------

        resume_json = (
            f"resume_{attempt}.json"
        )

        extract_docx(
            current_docx,
            resume_json
        )

        # ----------------------------------------------------
        # ATS
        # ----------------------------------------------------

        ats_json = (
            f"ats_{attempt}.json"
        )

        ats_analysis(
            resume_json,
            jd_path,
            ats_json
        )

        result = load_json(
            ats_json
        )

        score = int(
            result["ats_score"]
        )

        print(
            f"ATS = {score}"
        )

        # ----------------------------------------------------
        # Best resume tracking
        # ----------------------------------------------------

        if score > best_score:

            best_score = score
            best_resume = current_docx

            print(
                f"New BEST score: {best_score}"
            )

        # ----------------------------------------------------
        # Target reached
        # ----------------------------------------------------

        if score >= target_score:

            print(
                "TARGET REACHED"
            )

            break

        # ----------------------------------------------------
        # Rewrite
        # ----------------------------------------------------

        rewrite_json = (
            f"rewrite_{attempt}.json"
        )

        rewrite_resume(
            current_docx,
            resume_json,
            jd_path,
            rewrite_json
        )

        # ----------------------------------------------------
        # Apply
        # ----------------------------------------------------

        next_docx = (
            f"resume_optimized_{attempt}.docx"
        )

        apply_rewrite(
            current_docx,
            rewrite_json,
            next_docx
        )

        # ----------------------------------------------------
        # Continue with optimized resume
        # ----------------------------------------------------

        current_docx = next_docx

        # ----------------------------------------------------
        # Evaluate optimized resume immediately
        # ----------------------------------------------------

        optimized_resume_json = f"optimized_check_{attempt}.json"
        optimized_ats_json = f"optimized_ats_{attempt}.json"

        extract_docx(
            current_docx,
            optimized_resume_json
        )

        ats_analysis(
            optimized_resume_json,
            jd_path,
            optimized_ats_json
        )

        optimized_result = load_json(
            optimized_ats_json
        )

        optimized_score = int(
            optimized_result["ats_score"]
        )

        print(
            f"Optimized ATS = {optimized_score}"
        )

        if optimized_score > best_score:
            best_score = optimized_score
            best_resume = current_docx

            print(
                f"New BEST optimized score: {best_score}"
            )

    # --------------------------------------------------------
    # Final best resume
    # --------------------------------------------------------

    shutil.copy(
        best_resume,
        "best_resume.docx"
    )

    print(
        "\n"
        + "=" * 60
    )

    print(
        f"BEST ATS SCORE = {best_score}"
    )

    print(
        f"BEST RESUME = best_resume.docx"
    )

    print(
        "=" * 60
    )


# ============================================================
# MAIN
# ============================================================

def main():

    if len(sys.argv) < 2:

        print(
            """
Usage:

python3 resume_ai.py extract input.docx output.json

python3 resume_ai.py ats resume.json jd.txt ats.json

python3 resume_ai.py rewrite input.docx resume.json jd.txt rewritten.json

python3 resume_ai.py apply input.docx rewritten.json output.docx

python3 resume_ai.py optimize input.docx jd.txt
"""
        )

        sys.exit(1)

    command = sys.argv[1]

    # --------------------------------------------------------
    # EXTRACT
    # --------------------------------------------------------

    if command == "extract":

        extract_docx(
            sys.argv[2],
            sys.argv[3]
        )

    # --------------------------------------------------------
    # ATS
    # --------------------------------------------------------

    elif command == "ats":

        ats_analysis(
            sys.argv[2],
            sys.argv[3],
            sys.argv[4]
        )

    # --------------------------------------------------------
    # REWRITE
    # --------------------------------------------------------

    elif command == "rewrite":

        rewrite_resume(
            sys.argv[2],
            sys.argv[3],
            sys.argv[4],
            sys.argv[5]
        )

    # --------------------------------------------------------
    # APPLY
    # --------------------------------------------------------

    elif command == "apply":

        apply_rewrite(
            sys.argv[2],
            sys.argv[3],
            sys.argv[4]
        )

    # --------------------------------------------------------
    # OPTIMIZE
    # --------------------------------------------------------

    elif command == "optimize":

        optimize_resume(
            sys.argv[2],
            sys.argv[3]
        )

    # --------------------------------------------------------
    # UNKNOWN
    # --------------------------------------------------------

    else:

        raise RuntimeError(
            f"Unknown command: {command}"
        )


# ============================================================
# ENTRY POINT
# ============================================================

if __name__ == "__main__":

    main()
